from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse
from pydantic import BaseModel
import os
import uuid
from pathlib import Path
import shutil
from typing import Optional
import asyncio
import random
import warnings
from getpass import getpass
import uvicorn
import uvicorn

warnings.filterwarnings('ignore')

# === JOB SEEKER ENCOURAGEMENT MESSAGES ===
ENCOURAGEMENT_MESSAGES = [
    "You've got this - every step forward is progress toward your goals! 💪",
    "Your unique skills and experiences make you a valuable candidate! 🌟",
    "Keep pushing forward - the right opportunity is out there waiting for you! 🚀",
    "Remember, every 'no' brings you one step closer to your 'yes'! ✨",
    "Your determination and effort will pay off - stay confident! 💼",
    "You're building something amazing with each application and interview! 🔥",
    "Trust in your abilities - you have more to offer than you realize! 🎯",
    "The job search journey is tough, but you're tougher! 💯",
    "Your persistence is your superpower - keep going! ⚡",
    "Great things are coming your way - stay positive and focused! 🌈",
    "Every experience is preparing you for the perfect role! 📈",
    "You're not just looking for any job - you're finding your next chapter! 📚",
    "Your skills are valuable, your experience matters, and your future is bright! ☀️",
    "The right employer will recognize your worth - don't settle for less! 👑",
    "You're investing in your future with every effort you make! 🎊"
]

def get_encouragement():
    """Get a random encouragement message for job seekers"""
    return random.choice(ENCOURAGEMENT_MESSAGES)

# === RAG IMPORTS ===
try:
    from ibm_watsonx_ai.foundation_models import ModelInference
    from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
    from ibm_watsonx_ai.metanames import EmbedTextParamsMetaNames
    from langchain_ibm import WatsonxLLM, WatsonxEmbeddings
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import Chroma
    from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
    from langchain.chains import RetrievalQA
    RAG_AVAILABLE = True
    print("✅ RAG dependencies loaded successfully")
except ImportError as e:
    print(f"⚠️  RAG dependencies not available: {e}")
    print("📦 Install with: pip install langchain-ibm langchain-community chromadb ibm-watsonx-ai python-docx")
    RAG_AVAILABLE = False

# === CREDENTIAL SETUP ===
def setup_credentials():
    """Setup and validate IBM Watson credentials"""
    print("🔍 Checking IBM Watson credentials...")

    from dotenv import load_dotenv
    from pathlib import Path
    import os

    # Try environment variables first
    ibm_api_key = os.getenv('IBM_API_KEY')
    ibm_project_id = os.getenv('IBM_PROJECT_ID')
    
    if ibm_api_key and ibm_project_id:
        print("✅ Found credentials in environment variables")
        # Make sure they're set in os.environ for RAG functions
        os.environ['IBM_API_KEY'] = ibm_api_key
        os.environ['IBM_PROJECT_ID'] = ibm_project_id
        return ibm_api_key, ibm_project_id
    
    # Try loading from ~/.env specifically
    try:
        # Load from ~/.env specifically
        env_path = Path.home() / '.env'
        
        if env_path.exists():
            print(f"🔍 Loading from {env_path}")
            load_dotenv(env_path)
            
            ibm_api_key = os.getenv('IBM_API_KEY')
            ibm_project_id = os.getenv('IBM_PROJECT_ID')
            
            if ibm_api_key and ibm_project_id:
                print("✅ Found credentials from ~/.env file")
                # IMPORTANT: Set them in os.environ so RAG functions can access them
                os.environ['IBM_API_KEY'] = ibm_api_key
                os.environ['IBM_PROJECT_ID'] = ibm_project_id
                print(f"   API Key: {'*' * 20}...{ibm_api_key[-10:] if len(ibm_api_key) > 30 else 'present'}")
                print(f"   Project ID: {'*' * 20}...{ibm_project_id[-10:] if len(ibm_project_id) > 30 else 'present'}")
                return ibm_api_key, ibm_project_id
            else:
                print(f"⚠️  ~/.env file exists but missing credentials")
                print(f"   IBM_API_KEY: {'found' if ibm_api_key else 'missing'}")
                print(f"   IBM_PROJECT_ID: {'found' if ibm_project_id else 'missing'}")
        else:
            print(f"⚠️  ~/.env file not found at {env_path}")
            
    except ImportError:
        print("⚠️  python-dotenv not installed")
        print("   Install with: pip install python-dotenv")
    except Exception as e:
        print(f"⚠️  Error loading ~/.env: {e}")
    
    # If still missing, return None (will use demo mode)
    print("⚠️  IBM Watson credentials not found - using demo mode")
    return None, None

# === RAG FUNCTIONS ===
def get_llm():
    """Initialize the LLM with error handling"""
    if not RAG_AVAILABLE:
        raise ValueError("RAG dependencies not available")
    
    try:
        model_id = 'mistralai/mixtral-8x7b-instruct-v01'
        parameters = {
            GenParams.MAX_NEW_TOKENS: 256,
            GenParams.TEMPERATURE: 0.5,
        }
        
        watsonx_llm = WatsonxLLM(
            model_id=model_id,
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=parameters,
        )
        return watsonx_llm
    except Exception as e:
        print(f"❌ Error initializing LLM: {e}")
        raise

def watsonx_embedding():
    """Initialize embeddings with error handling"""
    if not RAG_AVAILABLE:
        raise ValueError("RAG dependencies not available")
    
    try:
        embed_params = {
            EmbedTextParamsMetaNames.TRUNCATE_INPUT_TOKENS: 3,
            EmbedTextParamsMetaNames.RETURN_OPTIONS: {"input_text": True},
        }
        
        watsonx_embedding = WatsonxEmbeddings(
            model_id="ibm/slate-125m-english-rtrvr",
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=embed_params,
        )
        return watsonx_embedding
    except Exception as e:
        print(f"❌ Error initializing embeddings: {e}")
        raise

def get_file_icon(file_extension):
    """Get appropriate emoji icon for file type"""
    icons = {
        '.pdf': '📄',
        '.doc': '📝',
        '.docx': '📝',
        '.txt': '📃'
    }
    return icons.get(file_extension.lower(), '📎')

def document_loader_universal(file_path):
    """Universal document loader supporting PDF, DOC, DOCX, and TXT files"""
    if not RAG_AVAILABLE:
        raise ValueError("RAG dependencies not available")
    
    try:
        # Get file extension and icon
        file_extension = Path(file_path).suffix.lower()
        file_icon = get_file_icon(file_extension)
        
        print(f"{file_icon} Loading {file_extension.upper()} file: {file_path}")
        
        # Route to appropriate loader based on file type
        if file_extension == '.pdf':
            loader = PyPDFLoader(file_path)
            
        elif file_extension in ['.doc', '.docx']:
            try:
                loader = Docx2txtLoader(file_path)
            except Exception as docx_error:
                print(f"⚠️  Docx2txtLoader failed: {docx_error}")
                try:
                    from docx import Document
                    from langchain.schema import Document as LangChainDocument
                    
                    if file_extension == '.docx':
                        doc = Document(file_path)
                        full_text = []
                        for paragraph in doc.paragraphs:
                            full_text.append(paragraph.text)
                        
                        content = '\n'.join(full_text)
                        if not content.strip():
                            raise ValueError("DOCX file appears to be empty")
                        
                        loaded_document = [LangChainDocument(
                            page_content=content,
                            metadata={'source': file_path, 'file_type': file_extension}
                        )]
                        
                        print(f"✅ Successfully loaded DOCX using fallback method")
                        return _finalize_document_loading(loaded_document, file_path, file_extension)
                    else:
                        raise ValueError(".DOC files require conversion to .DOCX format")
                        
                except ImportError:
                    raise ValueError("python-docx package required for DOCX files")
                except Exception as fallback_error:
                    raise ValueError(f"Failed to load {file_extension} file: {fallback_error}")
            
        elif file_extension == '.txt':
            encodings_to_try = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
            
            loader = None
            for encoding in encodings_to_try:
                try:
                    loader = TextLoader(file_path, encoding=encoding)
                    test_load = loader.load()
                    print(f"✅ Successfully detected text encoding: {encoding}")
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            if loader is None:
                raise ValueError(f"Could not determine text file encoding")
                
        else:
            supported_types = ['.pdf', '.doc', '.docx', '.txt']
            raise ValueError(f"Unsupported file type: {file_extension}. Supported types: {supported_types}")
        
        # Load the document
        loaded_document = loader.load()
        return _finalize_document_loading(loaded_document, file_path, file_extension)
        
    except Exception as e:
        error_msg = f"Error loading document: {str(e)}"
        print(f"❌ {error_msg}")
        raise ValueError(error_msg)

def _finalize_document_loading(loaded_document, file_path, file_extension):
    """Helper function to finalize document loading with validation and metadata"""
    if not loaded_document:
        raise ValueError(f"No content found in file")
    
    total_chars = sum(len(doc.page_content) for doc in loaded_document)
    if total_chars == 0:
        raise ValueError(f"File appears to be empty")
    
    file_icon = get_file_icon(file_extension)
    print(f"✅ Successfully loaded {len(loaded_document)} document section(s) with {total_chars:,} characters")
    
    # Enhanced metadata
    for i, doc in enumerate(loaded_document):
        doc.metadata.update({
            'file_type': file_extension,
            'source_file': os.path.basename(file_path),
            'file_size_bytes': os.path.getsize(file_path),
            'total_sections': len(loaded_document),
            'section_index': i,
            'character_count': len(doc.page_content),
            'file_icon': file_icon
        })
    
    return loaded_document

def text_splitter(data):
    """Split text into chunks with improved configuration"""
    if not RAG_AVAILABLE:
        raise ValueError("RAG dependencies not available")
    
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_documents(data)
        
        if not chunks:
            raise ValueError("No chunks created from document")
        
        print(f"✅ Split document into {len(chunks)} chunks")
        return chunks
    except Exception as e:
        print(f"❌ Error splitting text: {e}")
        raise

def vector_database_persistant(chunks, file_id):
    """Create persistent vector database"""
    try:
        embedding_model = watsonx_embedding()
        
        # Create persistent storage directory
        persist_dir = Path("vector_db") / file_id
        persist_dir.mkdir(parents=True, exist_ok=True)
        
        vectordb = Chroma.from_documents(
            documents=chunks, 
            embedding=embedding_model,
            collection_metadata={"hnsw:space": "cosine"},
            persist_directory=str(persist_dir)  # ← This makes it persistent
        )
        
        # Explicitly persist
        vectordb.persist()
        
        print(f"✅ Created persistent vector database at {persist_dir}")
        return vectordb    
    except Exception as e:
        print(f"❌ Error creating persistant vector database: {e}")
        raise
    
def vector_database(chunks):
    """Create vector database with improved configuration"""
    if not RAG_AVAILABLE:
        raise ValueError("RAG dependencies not available")
    
    try:
        embedding_model = watsonx_embedding()
        
        vectordb = Chroma.from_documents(
            documents=chunks, 
            embedding=embedding_model,
            collection_metadata={"hnsw:space": "cosine"}
        )
        print("✅ Created vector database")
        return vectordb
    except Exception as e:
        print(f"❌ Error creating vector database: {e}")
        raise

def create_retriever(file_path):
    """Create retriever from file with improved configuration"""
    if not RAG_AVAILABLE:
        raise ValueError("RAG dependencies not available")
    
    try:
        splits = document_loader_universal(file_path)
        chunks = text_splitter(splits)
        vectordb = vector_database(chunks)
        
        max_chunks = len(chunks)
        k = min(4, max_chunks)
        
        retriever = vectordb.as_retriever(
            search_type="similarity",
            search_kwargs={"k": k}
        )
        print(f"✅ Created retriever with k={k} (max available: {max_chunks})")
        return retriever, splits[0].metadata if splits else {}
    except Exception as e:
        print(f"❌ Error creating retriever: {e}")
        raise

def retriever_qa(file_path, query):
    """Main QA function with comprehensive error handling"""
    if not RAG_AVAILABLE:
        # Fallback to demo response
        encouragement = get_encouragement()
        return f"🤖 Demo Mode: I would analyze your document '{os.path.basename(file_path)}' to answer: '{query}'\n\nTo enable real RAG processing, install: pip install langchain-ibm langchain-community chromadb ibm-watsonx-ai python-docx\n\n🌟 {encouragement}"
    
    try:
        if not query or not query.strip():
            return "Please enter a question."
        
        print(f"🔍 Processing query: {query}")
        print(f"📎 File: {file_path}")
        
        # Initialize components
        llm = get_llm()
        retriever_obj, file_metadata = create_retriever(file_path)
        
        # Create QA chain
        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever_obj,
            return_source_documents=False,
            verbose=False
        )
        
        # Use invoke method
        try:
            response = qa.invoke({"query": query})
            result = response.get('result', response) if isinstance(response, dict) else response
        except AttributeError:
            response = qa({"query": query})
            result = response['result']
        
        # Add file info to response
        file_icon = file_metadata.get('file_icon', '📎')
        source_file = file_metadata.get('source_file', 'document')
        file_type = file_metadata.get('file_type', '').upper()
        
        # Get encouragement message
        encouragement = get_encouragement()
        
        # Combine response
        response_with_source = f"{result}\n\n---\n💡 *Source: {file_icon} {source_file} ({file_type})*\n\n🌟 {encouragement}"
        
        print("✅ Generated response successfully")
        return response_with_source
        
    except Exception as e:
        error_msg = f"Processing Error: {str(e)}"
        print(f"❌ {error_msg}")
        encouragement = get_encouragement()
        return f"Sorry, I encountered an error while processing your request.\n\nError details: {str(e)}\n\nPlease check your file and try again.\n\n🌟 {encouragement}"

# === FASTAPI APP SETUP ===
# Create FastAPI app
app = FastAPI(title="RAG Document Q&A System", version="1.0.0")

# Create logging
import logging
import os

import logging
from pathlib import Path

import logging
from pathlib import Path

# Create directory
Path("/var/log/FastAPI").mkdir(parents=True, exist_ok=True)

# Configure the specific uvicorn access logger
access_logger = logging.getLogger("uvicorn.access")
access_logger.setLevel(logging.INFO)

handler = logging.FileHandler('/var/log/FastAPI/uvicorn.access')
handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))
access_logger.addHandler(handler)

# Serve static files (CSS, JS, HTML)
app.mount("/static", StaticFiles(directory="static"), name="static")

# Create uploads directory if it doesn't exist
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Store uploaded files info (in production, use a database)
uploaded_files = {}

# Global RAG state
rag_initialized = False

class QueryRequest(BaseModel):
    query: str
    file_id: str

class QueryResponse(BaseModel):
    answer: str
    file_name: str

@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the main HTML page"""
    try:
        with open("static/index.html", "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Please make sure index.html is in the static/ directory</h1>")

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """Handle file upload and RAG processing"""
    
    # Validate file type
    allowed_types = {
        "application/pdf": ".pdf",
        "application/msword": ".doc", 
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt"
    }
    
    # Check file extension as fallback
    file_extension = Path(file.filename).suffix.lower()
    if file.content_type not in allowed_types and file_extension not in ['.pdf', '.doc', '.docx', '.txt']:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed types: PDF, DOC, DOCX, TXT"
        )
    
    # Generate unique file ID
    file_id = str(uuid.uuid4())
    
    # Determine file extension
    if file.content_type in allowed_types:
        file_extension = allowed_types[file.content_type]
    else:
        file_extension = Path(file.filename).suffix.lower()
    
    # Create file path
    file_path = UPLOAD_DIR / f"{file_id}{file_extension}"
    
    try:
        # Save file to disk
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Store file info
        uploaded_files[file_id] = {
            "original_name": file.filename,
            "file_path": str(file_path),
            "content_type": file.content_type,
            "file_extension": file_extension,
            "processed": False,
            "processing_error": None
        }
        
        # Process document for RAG - but ensure credentials are loaded first
        try:
            if RAG_AVAILABLE:
                # CRITICAL FIX: Always ensure credentials are loaded before RAG processing
                if not (os.getenv('IBM_API_KEY') and os.getenv('IBM_PROJECT_ID')):
                    print("🔍 Credentials not in environment, loading them now...")
                    api_key, project_id = setup_credentials()
                    if not (api_key and project_id):
                        print("⚠️  No credentials available, using demo mode")
                        # Don't modify rag_initialized here since it's global
                    else:
                        # Credentials are set by setup_credentials()
                        global rag_initialized
                        rag_initialized = True
                
                if rag_initialized and os.getenv('IBM_API_KEY') and os.getenv('IBM_PROJECT_ID'):
                    # Test if we can create retriever (this validates the document)
                    _, file_metadata = create_retriever(str(file_path))
                    uploaded_files[file_id]["processed"] = True
                    uploaded_files[file_id]["file_metadata"] = file_metadata
                    print(f"✅ Successfully processed {file.filename} for RAG")
                else:
                    uploaded_files[file_id]["processed"] = True  # Mark as processed for demo mode
                    print(f"📋 File uploaded in demo mode: {file.filename}")
            else:
                uploaded_files[file_id]["processed"] = True  # Mark as processed for demo mode
                print(f"📋 File uploaded in demo mode (RAG dependencies not available): {file.filename}")
                
        except Exception as processing_error:
            print(f"⚠️  RAG processing failed: {processing_error}")
            uploaded_files[file_id]["processed"] = True  # Still allow querying
            uploaded_files[file_id]["processing_error"] = str(processing_error)
        
        file_icon = get_file_icon(file_extension)
        
        return JSONResponse({
            "status": "success",
            "message": f"File '{file.filename}' uploaded successfully!",
            "file_id": file_id,
            "file_name": file.filename,
            "file_icon": file_icon,
            "rag_available": RAG_AVAILABLE and rag_initialized,
            "processing_error": uploaded_files[file_id].get("processing_error")
        })
        
    except Exception as e:
        # Clean up file if something went wrong
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")

@app.post("/api/query", response_model=QueryResponse)
async def query_document(request: QueryRequest):
    """Process queries against uploaded documents using real RAG"""
    
    if request.file_id not in uploaded_files:
        raise HTTPException(status_code=404, detail="File not found")
    
    file_info = uploaded_files[request.file_id]
    
    if not file_info["processed"]:
        raise HTTPException(status_code=400, detail="File is still being processed")
    
    try:
        # CRITICAL FIX: Ensure credentials are loaded before any RAG processing
        if RAG_AVAILABLE:
            if not (os.getenv('IBM_API_KEY') and os.getenv('IBM_PROJECT_ID')):
                print("🔍 Credentials not in environment for query, loading them now...")
                api_key, project_id = setup_credentials()
                if not (api_key and project_id):
                    # Fall back to demo mode
                    encouragement = get_encouragement()
                    demo_answer = f"🤖 Demo Mode: I would analyze your document '{file_info['original_name']}' to answer: '{request.query}'\n\nIBM Watson credentials not available.\n\n🌟 {encouragement}"
                    
                    return QueryResponse(
                        answer=demo_answer,
                        file_name=file_info["original_name"]
                    )
        
        # Use real RAG processing
        file_path = file_info["file_path"]
        answer = retriever_qa(file_path, request.query)
        
        return QueryResponse(
            answer=answer,
            file_name=file_info["original_name"]
        )
        
    except Exception as e:
        # Fallback to demo response with encouragement
        encouragement = get_encouragement()
        demo_answer = f"🤖 I would analyze your document '{file_info['original_name']}' to answer: '{request.query}'\n\n[Real RAG processing failed: {str(e)}]\n\n🌟 {encouragement}"
        
        return QueryResponse(
            answer=demo_answer,
            file_name=file_info["original_name"]
        )

@app.get("/api/files")
async def list_files():
    """List all uploaded files"""
    return {
        "files": [
            {
                "file_id": file_id,
                "name": info["original_name"],
                "processed": info["processed"]
            }
            for file_id, info in uploaded_files.items()
        ]
    }

@app.delete("/api/files/{file_id}")
async def delete_file(file_id: str):
    """Delete an uploaded file"""
    if file_id not in uploaded_files:
        raise HTTPException(status_code=404, detail="File not found")
    
    file_info = uploaded_files[file_id]
    
    # Delete file from disk
    file_path = Path(file_info["file_path"])
    if file_path.exists():
        file_path.unlink()
    
    # Remove from memory
    del uploaded_files[file_id]
    
    return {"status": "success", "message": "File deleted successfully"}

@app.get("/api/debug")
async def debug_info():
    """Debug endpoint to check system status"""
    return {
        "uploads_dir_exists": UPLOAD_DIR.exists(),
        "uploads_dir_writable": os.access(UPLOAD_DIR, os.W_OK),
        "static_dir_exists": Path("static").exists(),
        "index_html_exists": Path("static/index.html").exists(),
        "rag_available": RAG_AVAILABLE,
        "rag_initialized": rag_initialized,
        "env_vars": {
            "IBM_API_KEY": "***" if os.getenv('IBM_API_KEY') else None,
            "IBM_PROJECT_ID": "***" if os.getenv('IBM_PROJECT_ID') else None
        }
    }

@app.get("/api/status")
async def get_system_status():
    """Get system status including RAG availability"""
    return {
        "rag_available": RAG_AVAILABLE,
        "rag_initialized": rag_initialized,
        "credentials_configured": bool(os.getenv('IBM_API_KEY') and os.getenv('IBM_PROJECT_ID')),
        "uploaded_files_count": len(uploaded_files)
    }

@app.post("/api/initialize")
async def initialize_rag():
    """Initialize RAG system with credentials"""
    global rag_initialized
    
    if not RAG_AVAILABLE:
        raise HTTPException(status_code=400, detail="RAG dependencies not available")
    
    try:
        # Setup credentials
        api_key, project_id = setup_credentials()
        
        if not api_key or not project_id:
            raise HTTPException(
                status_code=400, 
                detail="IBM Watson credentials not found. Set IBM_API_KEY and IBM_PROJECT_ID environment variables."
            )
        
        # Credentials are already set in os.environ by setup_credentials()
        
        # Test the connection
        test_llm = get_llm()
        test_embedding = watsonx_embedding()
        
        rag_initialized = True
        return {"status": "success", "message": "RAG system initialized successfully"}
        
    except Exception as e:
        rag_initialized = False
        raise HTTPException(status_code=500, detail=f"Failed to initialize RAG system: {str(e)}")

# Legacy functions for compatibility
async def simulate_document_processing(file_id: str):
    """Legacy function - kept for compatibility"""
    await asyncio.sleep(0.1)  # Minimal delay
    uploaded_files[file_id]["processed"] = True

async def simulate_rag_query(query: str, file_info: dict) -> str:
    """Legacy function - kept for compatibility"""
    encouragement = get_encouragement()
    return f"Demo response for query: '{query}' on file '{file_info['original_name']}'\n\n🌟 {encouragement}"

if __name__ == "__main__":
    print("🚀 Starting RAG Document Q&A System...")
    print("📋 Supported formats: PDF, DOC, DOCX, TXT")
    
    # Initialize RAG if credentials are available
    try:
        if RAG_AVAILABLE:
            api_key, project_id = setup_credentials()
            if api_key and project_id:
                # Credentials are already set in os.environ by setup_credentials()
                rag_initialized = True
                print("✅ RAG system initialized with IBM Watson")
            else:
                print("⚠️  Running in demo mode - set IBM_API_KEY and IBM_PROJECT_ID for full functionality")
        else:
            print("⚠️  RAG dependencies not available - running in demo mode")
    except Exception as e:
        print(f"⚠️  RAG initialization failed: {e}")
    
    print("\n📁 Make sure to create:")
    print("1. 'static' directory with index.html and styles.css")
    print("2. Install dependencies:")
    print("   pip install fastapi uvicorn python-multipart")
    print("   pip install langchain-ibm langchain-community chromadb ibm-watsonx-ai python-docx  # For RAG")
    
    uvicorn.run(app, host="0.0.0.0", port=7862, reload=False)

