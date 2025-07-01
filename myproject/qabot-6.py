# -*- coding: utf-8 -*-
"""
RAG Application with improved error handling and credential validation
Now supports PDF, DOC, DOCX, and TXT files
"""

import os
from getpass import getpass
import warnings
warnings.filterwarnings('ignore')

# === CREDENTIAL SETUP WITH VALIDATION ===
def setup_credentials():
    """Setup and validate IBM Watson credentials - simplified sequential steps"""
    import os
    from getpass import getpass
    
    # Step 1: Try os.getenv first (checks if already in environment)
    print("🔍 Step 1: Checking environment variables...")
    ibm_api_key = os.getenv('IBM_API_KEY')
    ibm_project_id = os.getenv('IBM_PROJECT_ID')
    watsonx_apikey = os.getenv('WATSONX_APIKEY')
    
    if ibm_api_key and ibm_project_id:
        print("✅ Found credentials in environment variables")
        if not watsonx_apikey:
            watsonx_apikey = ibm_api_key
    else:
        # Step 2: Try dotenv to load .env file
        print("🔍 Step 2: Trying to load .env file...")
        try:
            from dotenv import load_dotenv
            from pathlib import Path
            
            # Try home directory first, then find_dotenv
            env_path = Path.home() / '.env'
            if env_path.exists():
                load_dotenv(env_path)
                print(f"📄 Loaded .env from {env_path}")
            else:
                from dotenv import find_dotenv
                env_file = find_dotenv()
                if env_file:
                    load_dotenv(env_file)
                    print(f"📄 Loaded .env from {env_file}")
            
            # Check environment again after loading .env
            ibm_api_key = os.getenv('IBM_API_KEY')
            ibm_project_id = os.getenv('IBM_PROJECT_ID')
            watsonx_apikey = os.getenv('WATSONX_APIKEY')
            
            if ibm_api_key and ibm_project_id:
                print("✅ Found credentials from .env file")
                if not watsonx_apikey:
                    watsonx_apikey = ibm_api_key
            else:
                print("⚠️  .env file loaded but missing required credentials")
                
        except ImportError:
            print("⚠️  python-dotenv not installed")
        except Exception as e:
            print(f"⚠️  Error loading .env: {e}")
    
    # Step 4: Manual input if still missing
    if not ibm_api_key or not ibm_project_id:
        print("🔍 Step 4: Manual input required...")
        ibm_api_key = getpass('IBM API Key: ')
        ibm_project_id = getpass('IBM Project ID: ')
        watsonx_apikey = ibm_api_key
    
    # Step 5: Validate and set
    if not ibm_api_key or len(ibm_api_key) < 40:
        raise ValueError("IBM API Key appears invalid (too short)")
    if not ibm_project_id or len(ibm_project_id) < 30:
        raise ValueError("IBM Project ID appears invalid (too short)")
    
    # Set environment variables
    os.environ['IBM_API_KEY'] = ibm_api_key
    os.environ['IBM_PROJECT_ID'] = ibm_project_id
    os.environ['WATSONX_APIKEY'] = watsonx_apikey
    
    print("✅ Credentials configured successfully")
    return ibm_api_key, ibm_project_id, watsonx_apikey

def setup_credentials1():
    """Setup and validate IBM Watson credentials"""
    try:
        # Try to get from userdata first
        ibm_api_key = userdata.get('IBM_API_KEY')
        ibm_project_id = userdata.get('IBM_PROJECT_ID')
        watsonx_apikey = userdata.get('WATSONX_APIKEY')
        
        # If not found, prompt user
        if not ibm_api_key or not ibm_project_id:
            print("⚠️  Credentials not found in userdata. Please enter manually:")
            ibm_api_key = getpass('IBM API Key: ')
            ibm_project_id = getpass('IBM Project ID: ')
            watsonx_apikey = ibm_api_key  # Often the same
        
        # Validate credentials format
        if not ibm_api_key or len(ibm_api_key) < 40:
            raise ValueError("IBM API Key appears invalid (too short)")
        if not ibm_project_id or len(ibm_project_id) < 30:
            raise ValueError("IBM Project ID appears invalid (too short)")
            
        # Set environment variables
        os.environ['IBM_API_KEY'] = ibm_api_key
        os.environ['IBM_PROJECT_ID'] = ibm_project_id
        os.environ['WATSONX_APIKEY'] = watsonx_apikey
        
        print("✅ Credentials configured successfully")
        return ibm_api_key, ibm_project_id, watsonx_apikey
        
    except Exception as e:
        print(f"❌ Error setting up credentials: {e}")
        raise

# === TEST CREDENTIALS FUNCTION ===
def test_credentials():
    """Test if credentials work by making a simple API call"""
    try:
        from ibm_watsonx_ai.foundation_models import ModelInference
        from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
        
        # Try to create a simple model instance
        model = ModelInference(
            model_id='mistralai/mixtral-8x7b-instruct-v01',
            params={GenParams.MAX_NEW_TOKENS: 10},
            credentials={
                "url": "https://us-south.ml.cloud.ibm.com",
                "apikey": os.environ['IBM_API_KEY']
            },
            project_id=os.environ['IBM_PROJECT_ID']
        )
        
        # Try a simple generation
        response = model.generate_text("Hello")
        print("✅ Credentials test successful!")
        return True
        
    except Exception as e:
        print(f"❌ Credentials test failed: {e}")
        return False

# === IMPORTS ===
from ibm_watsonx_ai.foundation_models import ModelInference
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
from ibm_watsonx_ai.metanames import EmbedTextParamsMetaNames
from langchain_ibm import WatsonxLLM, WatsonxEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain.chains import RetrievalQA
from langchain.schema import Document
import gradio as gr

# === ENHANCED DOCUMENT LOADING FUNCTIONS ===

def load_txt_file(file_path):
    """Load a TXT file and return as Document objects"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Create a Document object similar to what PyPDFLoader returns
        doc = Document(page_content=content, metadata={"source": file_path})
        print(f"✅ Loaded TXT file: {len(content)} characters")
        return [doc]  # Return as list to match PDF loader format
        
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
            doc = Document(page_content=content, metadata={"source": file_path})
            print(f"✅ Loaded TXT file with latin-1 encoding: {len(content)} characters")
            return [doc]
        except Exception as e:
            print(f"❌ Error loading TXT file with fallback encoding: {e}")
            raise
    except Exception as e:
        print(f"❌ Error loading TXT file: {e}")
        raise

def load_docx_file(file_path):
    """Load a DOCX file and return as Document objects"""
    try:
        from docx import Document as DocxDocument
        
        # Load the DOCX file
        doc = DocxDocument(file_path)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text)
        
        # Join all paragraphs with newlines
        full_text = '\n'.join(text_content)
        
        # Create Document object
        document = Document(page_content=full_text, metadata={"source": file_path})
        print(f"✅ Loaded DOCX file: {len(doc.paragraphs)} paragraphs, {len(full_text)} characters")
        return [document]
        
    except ImportError:
        error_msg = "python-docx library not installed. Install with: pip install python-docx"
        print(f"❌ {error_msg}")
        raise ImportError(error_msg)
    except Exception as e:
        print(f"❌ Error loading DOCX file: {e}")
        raise

def load_doc_file(file_path):
    """Load a DOC file and return as Document objects"""
    try:
        import subprocess
        import tempfile
        import os
        
        # Try using python-docx2txt for DOC files
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            document = Document(page_content=text, metadata={"source": file_path})
            print(f"✅ Loaded DOC file using docx2txt: {len(text)} characters")
            return [document]
        except ImportError:
            print("⚠️  docx2txt not available, trying alternative method...")
        
        # Alternative: Try using python-mammoth (better for complex formatting)
        try:
            import mammoth
            
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
                
            document = Document(page_content=text, metadata={"source": file_path})
            print(f"✅ Loaded DOC file using mammoth: {len(text)} characters")
            return [document]
            
        except ImportError:
            error_msg = ("No suitable library found for DOC files. Install one of:\n"
                        "- pip install docx2txt\n"
                        "- pip install mammoth")
            print(f"❌ {error_msg}")
            raise ImportError(error_msg)
            
    except Exception as e:
        print(f"❌ Error loading DOC file: {e}")
        raise

def get_file_extension(file_path):
    """Get file extension in lowercase"""
    return os.path.splitext(file_path)[1].lower()

def document_loader(file):
    """Enhanced document loader that handles PDF, DOC, DOCX, and TXT files"""
    try:
        file_path = file.name
        file_ext = get_file_extension(file_path)
        
        print(f"📄 Loading file: {file_path} (type: {file_ext})")
        
        # Route to appropriate loader based on file extension
        if file_ext == '.pdf':
            loader = PyPDFLoader(file_path)
            loaded_document = loader.load()
            print(f"✅ Loaded PDF: {len(loaded_document)} pages")
            
        elif file_ext == '.txt':
            loaded_document = load_txt_file(file_path)
            
        elif file_ext == '.docx':
            loaded_document = load_docx_file(file_path)
            
        elif file_ext == '.doc':
            loaded_document = load_doc_file(file_path)
            
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: .pdf, .doc, .docx, .txt")
        
        return loaded_document
        
    except Exception as e:
        print(f"❌ Error in document_loader: {e}")
        raise

# === CORE FUNCTIONS WITH ERROR HANDLING (unchanged) ===

def get_llm():
    """Initialize the LLM with error handling"""
    try:
        model_id = 'mistralai/mixtral-8x7b-instruct-v01'
        parameters = {
            GenParams.MAX_NEW_TOKENS: 256,
            GenParams.TEMPERATURE: 0.5,
        }
        
        watsonx_llm = WatsonxLLM(
            model_id=model_id,
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=parameters,
        )
        return watsonx_llm
    except Exception as e:
        print(f"❌ Error initializing LLM: {e}")
        raise

def watsonx_embedding():
    """Initialize embeddings with error handling"""
    try:
        embed_params = {
            EmbedTextParamsMetaNames.TRUNCATE_INPUT_TOKENS: 3,
            EmbedTextParamsMetaNames.RETURN_OPTIONS: {"input_text": True},
        }
        
        watsonx_embedding = WatsonxEmbeddings(
            model_id="ibm/slate-125m-english-rtrvr",
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],  # Use apikey instead of params
            params=embed_params,
        )
        return watsonx_embedding
    except Exception as e:
        print(f"❌ Error initializing embeddings: {e}")
        raise

def text_splitter(data):
    """Split text into chunks"""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=50,
            length_function=len,
        )
        chunks = text_splitter.split_documents(data)
        print(f"✅ Split document into {len(chunks)} chunks")
        return chunks
    except Exception as e:
        print(f"❌ Error splitting text: {e}")
        raise

def vector_database(chunks):
    """Create vector database"""
    try:
        embedding_model = watsonx_embedding()
        vectordb = Chroma.from_documents(chunks, embedding_model)
        print("✅ Created vector database")
        return vectordb
    except Exception as e:
        print(f"❌ Error creating vector database: {e}")
        raise

def retriever(file):
    """Create retriever from file"""
    try:
        splits = document_loader(file)
        chunks = text_splitter(splits)
        vectordb = vector_database(chunks)
        retriever = vectordb.as_retriever()
        print("✅ Created retriever")
        return retriever
    except Exception as e:
        print(f"❌ Error creating retriever: {e}")
        raise

def retriever_qa(file, query):
    """Main QA function with comprehensive error handling"""
    try:
        print(f"🔍 Processing query: {query}")
        print(f"📄 File: {file}")
        
        # Initialize components
        llm = get_llm()
        retriever_obj = retriever(file)
        
        # Create QA chain
        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever_obj,
            return_source_documents=False
        )
        
        # Get response
        response = qa(query)
        print("✅ Generated response successfully")
        return response['result']
        
    except Exception as e:
        error_msg = f"❌ Error processing query: {str(e)}"
        print(error_msg)
        return f"Sorry, I encountered an error: {str(e)}\n\nPlease check your credentials and try again."

# === MAIN SETUP ===
def main():
    """Main setup function"""
    print("🚀 Setting up Enhanced RAG Application...")
    
    # Setup credentials
    setup_credentials()
    
    # Test credentials
    if not test_credentials():
        print("❌ Credential test failed. Please check your IBM Watson credentials.")
        return None
    
    # Create Gradio interface with updated file types
    rag_application = gr.Interface(
        fn=retriever_qa,
        allow_flagging="never",
        inputs=[
            gr.File(
                label="Upload Document", 
                file_count="single", 
                file_types=['.pdf', '.doc', '.docx', '.txt'], 
                type="filepath"
            ),
            gr.Textbox(
                label="Input Query", 
                lines=2, 
                placeholder="Type your question here..."
            )
        ],
        outputs=gr.Textbox(label="Output"),
        title="Enhanced RAG Chatbot",
        description="Upload a document (PDF, DOC, DOCX, or TXT) and ask any question. The chatbot will try to answer using the provided document."
    )
    
    print("✅ Enhanced RAG Application ready!")
    print("📄 Supported file formats: PDF, DOC, DOCX, TXT")
    return rag_application

# === LAUNCH ===
if __name__ == "__main__":
    # Setup and launch
    app = main()
    if app:
        app.launch(server_name="0.0.0.0", server_port=7860, share=False, debug=False)
    else:
        print("❌ Failed to initialize application")
