# -*- coding: utf-8 -*-
"""
RAG Application with improved error handling, credential validation, and LangChain updates
Fixed for Gradio compatibility issues
"""

import os
from getpass import getpass
import warnings
warnings.filterwarnings('ignore')

# === CREDENTIAL SETUP WITH VALIDATION ===
def setup_credentials():
    """Setup and validate IBM Watson credentials - simplified sequential steps"""
    import os
    from getpass import getpass
    
    # Step 1: Try os.getenv first (checks if already in environment)
    print("🔍 Step 1: Checking environment variables...")
    ibm_api_key = os.getenv('IBM_API_KEY')
    ibm_project_id = os.getenv('IBM_PROJECT_ID')
    watsonx_apikey = os.getenv('WATSONX_APIKEY')
    
    if ibm_api_key and ibm_project_id:
        print("✅ Found credentials in environment variables")
        if not watsonx_apikey:
            watsonx_apikey = ibm_api_key
    else:
        # Step 2: Try dotenv to load .env file
        print("🔍 Step 2: Trying to load .env file...")
        try:
            from dotenv import load_dotenv
            from pathlib import Path
            
            # Try home directory first, then find_dotenv
            env_path = Path.home() / '.env'
            if env_path.exists():
                load_dotenv(env_path)
                print(f"📄 Loaded .env from {env_path}")
            else:
                from dotenv import find_dotenv
                env_file = find_dotenv()
                if env_file:
                    load_dotenv(env_file)
                    print(f"📄 Loaded .env from {env_file}")
            
            # Check environment again after loading .env
            ibm_api_key = os.getenv('IBM_API_KEY')
            ibm_project_id = os.getenv('IBM_PROJECT_ID')
            watsonx_apikey = os.getenv('WATSONX_APIKEY')
            
            if ibm_api_key and ibm_project_id:
                print("✅ Found credentials from .env file")
                if not watsonx_apikey:
                    watsonx_apikey = ibm_api_key
            else:
                print("⚠️  .env file loaded but missing required credentials")
                
        except ImportError:
            print("⚠️  python-dotenv not installed")
        except Exception as e:
            print(f"⚠️  Error loading .env: {e}")
    
    # Step 3: Manual input if still missing
    if not ibm_api_key or not ibm_project_id:
        print("🔍 Step 3: Manual input required...")
        ibm_api_key = getpass('IBM API Key: ')
        ibm_project_id = getpass('IBM Project ID: ')
        watsonx_apikey = ibm_api_key
    
    # Step 4: Validate and set
    if not ibm_api_key or len(ibm_api_key) < 40:
        raise ValueError("IBM API Key appears invalid (too short)")
    if not ibm_project_id or len(ibm_project_id) < 30:
        raise ValueError("IBM Project ID appears invalid (too short)")
    
    # Set environment variables
    os.environ['IBM_API_KEY'] = ibm_api_key
    os.environ['IBM_PROJECT_ID'] = ibm_project_id
    os.environ['WATSONX_APIKEY'] = watsonx_apikey
    
    print("✅ Credentials configured successfully")
    return ibm_api_key, ibm_project_id, watsonx_apikey

# === TEST CREDENTIALS FUNCTION ===
def test_credentials():
    """Test if credentials work by making a simple API call"""
    try:
        from ibm_watsonx_ai.foundation_models import ModelInference
        from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
        
        # Try to create a simple model instance
        model = ModelInference(
            model_id='mistralai/mixtral-8x7b-instruct-v01',
            params={GenParams.MAX_NEW_TOKENS: 10},
            credentials={
                "url": "https://us-south.ml.cloud.ibm.com",
                "apikey": os.environ['IBM_API_KEY']
            },
            project_id=os.environ['IBM_PROJECT_ID']
        )
        
        # Try a simple generation
        response = model.generate_text("Hello")
        print("✅ Credentials test successful!")
        return True
        
    except Exception as e:
        print(f"❌ Credentials test failed: {e}")
        return False

# === IMPORTS ===
from ibm_watsonx_ai.foundation_models import ModelInference
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
from ibm_watsonx_ai.metanames import EmbedTextParamsMetaNames
from langchain_ibm import WatsonxLLM, WatsonxEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain.chains import RetrievalQA
import gradio as gr

# === CORE FUNCTIONS WITH ERROR HANDLING ===

def get_llm():
    """Initialize the LLM with error handling"""
    try:
        model_id = 'mistralai/mixtral-8x7b-instruct-v01'
        parameters = {
            GenParams.MAX_NEW_TOKENS: 256,
            GenParams.TEMPERATURE: 0.5,
        }
        
        watsonx_llm = WatsonxLLM(
            model_id=model_id,
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=parameters,
        )
        return watsonx_llm
    except Exception as e:
        print(f"❌ Error initializing LLM: {e}")
        raise

def watsonx_embedding():
    """Initialize embeddings with error handling"""
    try:
        embed_params = {
            EmbedTextParamsMetaNames.TRUNCATE_INPUT_TOKENS: 3,
            EmbedTextParamsMetaNames.RETURN_OPTIONS: {"input_text": True},
        }
        
        watsonx_embedding = WatsonxEmbeddings(
            model_id="ibm/slate-125m-english-rtrvr",
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=embed_params,
        )
        return watsonx_embedding
    except Exception as e:
        print(f"❌ Error initializing embeddings: {e}")
        raise

def document_loader_advanced(file):
    """
    Enhanced document loader with support for multiple document formats
    
    Supports: PDF, DOCX, DOC, TXT files with automatic encoding detection
    """
    import os
    from pathlib import Path
    from langchain_community.document_loaders import PyPDFLoader, TextLoader
    
    try:
        # Debug logging
        print(f"🔍 Debug: Processing file input: {type(file)}")
        
        # Validate file input with more detailed checks
        if not file:
            raise ValueError("No file provided")
        
        # Handle different file input types
        if hasattr(file, 'name'):
            file_path = file.name
        elif isinstance(file, str):
            file_path = file
        else:
            raise ValueError(f"Invalid file type: {type(file)}. Expected file object or string path.")
        
        print(f"🔍 Debug: File path: {file_path}")
        
        if not os.path.exists(file_path):
            raise ValueError(f"File does not exist: {file_path}")
        
        # Get file extension
        file_extension = Path(file_path).suffix.lower()
        print(f"🔍 Debug: File extension: {file_extension}")
        
        # Route to appropriate loader
        if file_extension == '.pdf':
            print(f"📄 Loading PDF: {file_path}")
            try:
                loader = PyPDFLoader(file_path)
            except Exception as e:
                raise ValueError(f"Error creating PDF loader: {str(e)}")
            
        elif file_extension == '.docx':
            print(f"📝 Loading DOCX: {file_path}")
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                loader = Docx2txtLoader(file_path)
            except ImportError:
                raise ValueError("DOCX support requires 'docx2txt' package. Install with: pip install docx2txt")
            except Exception as e:
                raise ValueError(f"Error creating DOCX loader: {str(e)}")
            
        elif file_extension == '.doc':
            print(f"📝 Loading DOC: {file_path}")
            # First try unstructured
            try:
                from langchain_community.document_loaders import UnstructuredWordDocumentLoader
                loader = UnstructuredWordDocumentLoader(file_path)
                print("✅ Using UnstructuredWordDocumentLoader for DOC file")
            except ImportError:
                print("⚠️  UnstructuredWordDocumentLoader not available, trying python-docx fallback...")
                # Fallback: try to use python-docx for .doc files (limited support)
                try:
                    import docx
                    from langchain_core.documents import Document
                    
                    print("⚠️  Using python-docx fallback for .doc file (may have limited compatibility)")
                    doc = docx.Document(file_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    
                    if not content.strip():
                        raise ValueError("No readable content found in DOC file")
                    
                    # Create a Document object manually
                    loaded_document = [Document(
                        page_content=content,
                        metadata={
                            'source': file_path,
                            'file_type': '.doc'
                        }
                    )]
                    
                    print(f"✅ Successfully loaded DOC file with {len(content)} characters")
                    
                    # Add enhanced metadata
                    for doc in loaded_document:
                        doc.metadata.update({
                            'file_type': file_extension,
                            'source_file': os.path.basename(file_path),
                            'file_size_bytes': os.path.getsize(file_path),
                            'total_documents': len(loaded_document),
                            'character_count': len(doc.page_content)
                        })
                    
                    return loaded_document
                    
                except ImportError:
                    raise ValueError("DOC support requires either 'unstructured' package (pip install unstructured) or 'python-docx' package (pip install python-docx)")
                except Exception as e:
                    raise ValueError(f"Error reading DOC file with python-docx: {str(e)}. Try converting to DOCX or PDF format.")
            except Exception as e:
                raise ValueError(f"Error creating DOC loader: {str(e)}")
            
        elif file_extension == '.txt':
            print(f"📃 Loading TXT: {file_path}")
            # Try different encodings for better compatibility
            encodings_to_try = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            loader = None
            last_error = None
            
            for encoding in encodings_to_try:
                try:
                    print(f"🔍 Trying encoding: {encoding}")
                    loader = TextLoader(file_path, encoding=encoding)
                    # Test if it can actually load
                    test_load = loader.load()
                    print(f"✅ Successfully detected encoding: {encoding}")
                    break
                except UnicodeDecodeError as e:
                    last_error = f"UnicodeDecodeError with {encoding}: {str(e)}"
                    continue
                except Exception as e:
                    last_error = f"Error with {encoding}: {str(e)}"
                    continue
            
            if loader is None:
                raise ValueError(f"Could not determine text file encoding. Tried: {encodings_to_try}. Last error: {last_error}")
                
        else:
            supported_types = ['.pdf', '.docx', '.doc', '.txt']
            raise ValueError(f"Unsupported file type: {file_extension}. Supported types: {supported_types}")
        
        # Load the document (for cases where loader was created)
        if 'loaded_document' not in locals():
            print("🔍 Loading document with loader...")
            try:
                loaded_document = loader.load()
            except Exception as e:
                raise ValueError(f"Error loading document with loader: {str(e)}")
        
        # Validate content
        if not loaded_document:
            raise ValueError(f"No content found in file: {file_path}")
        
        total_chars = sum(len(doc.page_content) for doc in loaded_document if doc.page_content)
        if total_chars == 0:
            raise ValueError(f"File appears to be empty or contains no readable text: {file_path}")
        
        print(f"✅ Successfully loaded {len(loaded_document)} document(s) with {total_chars} characters")
        
        # Enhanced metadata
        try:
            for doc in loaded_document:
                if doc.metadata is None:
                    doc.metadata = {}
                doc.metadata.update({
                    'file_type': file_extension,
                    'source_file': os.path.basename(file_path),
                    'file_size_bytes': os.path.getsize(file_path),
                    'total_documents': len(loaded_document),
                    'character_count': len(doc.page_content) if doc.page_content else 0
                })
        except Exception as e:
            print(f"⚠️  Warning: Could not add metadata: {str(e)}")
        
        return loaded_document
        
    except Exception as e:
        error_msg = f"Error loading document: {str(e)}"
        print(f"❌ {error_msg}")
        print(f"❌ Debug info - File: {file}, Type: {type(file)}")
        raise ValueError(error_msg)

# Required installations (add to your requirements or install manually):
"""
# Core document processing
pip install python-docx     # For .docx files
pip install docx2txt        # Alternative for .docx files
pip install langchain-community  # For document loaders

# Optional for better .doc support
pip install unstructured    # For .doc files (recommended)
pip install python-magic    # For better file type detection

# Alternative lighter approach (if unstructured is too heavy):
# Just python-docx can handle some .doc files as fallback
"""    
    
def text_splitter(data):
    """Split text into chunks with improved configuration"""
    try:
        if not data:
            raise ValueError("No document data provided")
        
        # Improved text splitter configuration
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,  # Increased overlap for better continuity
            length_function=len,
            separators=["\n\n", "\n", " ", ""]  # Better separation hierarchy
        )
        chunks = text_splitter.split_documents(data)
        
        if not chunks:
            raise ValueError("No chunks created from document")
        
        print(f"✅ Split document into {len(chunks)} chunks")
        return chunks
    except Exception as e:
        print(f"❌ Error splitting text: {e}")
        raise

def vector_database(chunks):
    """Create vector database with improved configuration"""
    try:
        if not chunks:
            raise ValueError("No chunks provided for vector database")
        
        embedding_model = watsonx_embedding()
        
        # Configure Chroma with improved settings
        vectordb = Chroma.from_documents(
            documents=chunks, 
            embedding=embedding_model,
            collection_metadata={"hnsw:space": "cosine"}  # Better similarity metric
        )
        print("✅ Created vector database")
        return vectordb
    except Exception as e:
        print(f"❌ Error creating vector database: {e}")
        raise

def create_retriever(file):
    """Create retriever from file with improved configuration"""
    try:
        splits = document_loader_advanced(file)
        chunks = text_splitter(splits)
        vectordb = vector_database(chunks)
        
        # Configure retriever with dynamic k based on chunk count
        # This fixes the "Number of requested results greater than elements" warning
        max_chunks = len(chunks)
        k = min(4, max_chunks)  # Use minimum of 4 or available chunks
        
        retriever = vectordb.as_retriever(
            search_type="similarity",
            search_kwargs={"k": k}  # Dynamic k value
        )
        print(f"✅ Created retriever with k={k} (max available: {max_chunks})")
        return retriever
    except Exception as e:
        print(f"❌ Error creating retriever: {e}")
        raise

def retriever_qa(file, query):
    """Main QA function with comprehensive error handling and modern LangChain usage"""
    try:
        print(f"🔍 Debug: retriever_qa called with file={type(file)}, query='{query}'")
        
        # Input validation
        if not file:
            return "Please upload a document file first."
        
        if not query or not query.strip():
            return "Please enter a question."
        
        print(f"🔍 Processing query: {query}")
        print(f"📄 File: {file.name if hasattr(file, 'name') else file}")
        
        # Initialize components with error handling
        print("🔍 Debug: Initializing LLM...")
        try:
            llm = get_llm()
            print("✅ LLM initialized successfully")
        except Exception as e:
            return f"Error initializing LLM: {str(e)}"
        
        print("🔍 Debug: Creating retriever...")
        try:
            retriever_obj = create_retriever(file)
            print("✅ Retriever created successfully")
        except Exception as e:
            return f"Error creating retriever: {str(e)}"
        
        # Create QA chain
        print("🔍 Debug: Creating QA chain...")
        try:
            qa = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=retriever_obj,
                return_source_documents=False,
                verbose=False  # Reduce verbose output
            )
            print("✅ QA chain created successfully")
        except Exception as e:
            return f"Error creating QA chain: {str(e)}"
        
        # Generate response
        print("🔍 Debug: Generating response...")
        try:
            # Use invoke instead of deprecated __call__ method
            response = qa.invoke({"query": query})
            result = response.get('result', response) if isinstance(response, dict) else response
            print("✅ Generated response successfully")
            return result
        except AttributeError:
            try:
                # Fallback for older LangChain versions
                response = qa({"query": query})
                result = response['result']
                print("✅ Generated response successfully (fallback method)")
                return result
            except Exception as e:
                return f"Error generating response (fallback): {str(e)}"
        except Exception as e:
            return f"Error generating response: {str(e)}"
        
    except ValueError as ve:
        error_msg = f"Input Error: {str(ve)}"
        print(f"❌ {error_msg}")
        return error_msg
        
    except Exception as e:
        error_msg = f"Processing Error: {str(e)}"
        print(f"❌ {error_msg}")
        return f"Sorry, I encountered an error while processing your request.\n\nError details: {str(e)}\n\nPlease check your file and try again."

# === SIMPLIFIED GRADIO INTERFACE ===
def create_gradio_interface():
    """Create a simplified Gradio interface to avoid schema issues"""
    
    def process_query(file, query, history):
        """Process query and maintain chat history with enhanced error handling"""
        try:
            print(f"🔍 Debug: process_query called with file={type(file)}, query='{query}'")
            
            if not file:
                response = "⚠️ Please upload a document file first (.pdf, .docx, .doc, or .txt)."
            elif not query or not query.strip():
                response = "⚠️ Please enter a question."
            else:
                print(f"🔍 Debug: Calling retriever_qa...")
                response = retriever_qa(file, query)
                print(f"🔍 Debug: retriever_qa returned: {type(response)}")
            
            # Add to history
            if history is None:
                history = []
            history.append([query, response])
            return history, ""
            
        except Exception as e:
            error_msg = f"Error in process_query: {str(e)}"
            print(f"❌ {error_msg}")
            if history is None:
                history = []
            history.append([query, f"Sorry, an error occurred: {error_msg}"])
            return history, ""
    
    def clear_chat():
        """Clear chat history"""
        return []
    
    # Create the interface with simpler components
    with gr.Blocks(title="RAG Document Checker") as app:
        gr.Markdown("# 🤖 RAG Document Checker (WatsonxLLM)")
        gr.Markdown("Upload your document as a PDF, DOCX, DOC, or TXT file and ask questions about its content!")
        
        with gr.Row():
            with gr.Column(scale=1):
                # Simplified file upload
                file_upload = gr.File(
                    label="📄 Upload Document File", 
                    file_count="single", 
                    file_types=['.pdf', '.docx', '.doc', '.txt']
                )
                
                clear_btn = gr.Button("🗑️ Clear Chat")
                
            with gr.Column(scale=2):
                # Simplified chatbot
                chatbot = gr.Chatbot(label="Chat History", height=400)
                
                with gr.Row():
                    query_input = gr.Textbox(
                        label="Your Question",
                        placeholder="Ask a question about the uploaded document...",
                        lines=2,
                        scale=4
                    )
                    submit_btn = gr.Button("Send", scale=1)
        
        # Event handlers with error handling
        def safe_submit(file, query, history):
            try:
                return process_query(file, query, history)
            except Exception as e:
                print(f"❌ Error in safe_submit: {e}")
                if history is None:
                    history = []
                history.append([query, f"System error occurred: {str(e)}"])
                return history, ""
        
        submit_btn.click(
            safe_submit,
            inputs=[file_upload, query_input, chatbot],
            outputs=[chatbot, query_input]
        )
        
        query_input.submit(
            safe_submit,
            inputs=[file_upload, query_input, chatbot],
            outputs=[chatbot, query_input]
        )
        
        clear_btn.click(
            clear_chat,
            outputs=[chatbot]
        )
    
    return app

# === MAIN SETUP ===
def main():
    """Main setup function with enhanced error handling"""
    print("🚀 Setting up RAG Application...")
    
    try:
        # Setup credentials
        setup_credentials()
        
        # Test credentials
        if not test_credentials():
            print("❌ Credential test failed. Please check your IBM Watson credentials.")
            return None
        
        # Create Gradio interface
        app = create_gradio_interface()
        
        print("✅ RAG Application ready!")
        return app
        
    except Exception as e:
        print(f"❌ Failed to initialize application: {e}")
        return None

# === LAUNCH ===
if __name__ == "__main__":
    # Setup and launch
    app = main()
    if app:
        print("🌐 Launching application...")
        try:
            # Try different launch configurations to avoid the localhost issue
            import socket
            
            # Check if we're running in a remote environment
            hostname = socket.gethostname()
            
            if 'localhost' in hostname or hostname == 'ubuntu':
                # Remote environment - use share=True
                print("🔗 Detected remote environment, using share=True")
                app.launch(
                    share=True,  # This creates a public link
                    server_port=7860,
                    debug=False,
                    show_error=True
                )
            else:
                # Local environment
                print("🏠 Detected local environment")
                app.launch(
                    server_name="0.0.0.0",  # Use 127.0.0.1 instead of 0.0.0.0
                    server_port=7860,
                    share=False,
                    debug=False,
                    show_error=True
                )
                
        except Exception as launch_error:
            print(f"❌ Launch failed with standard settings: {launch_error}")
            print("🔄 Trying fallback launch with share=True...")
            try:
                app.launch(
                    share=True,
                    server_port=7860,
                    debug=False,
                    show_error=True
                )
            except Exception as fallback_error:
                print(f"❌ Fallback launch also failed: {fallback_error}")
                print("💡 Try running with: python your_script.py --share")
    else:
        print("❌ Failed to initialize application")

