# -*- coding: utf-8 -*-
"""
RAG Application with multi-format document support (PDF, DOC, DOCX, TXT)
Enhanced error handling, credential validation, and LangChain updates
"""

import os
from getpass import getpass
import warnings
warnings.filterwarnings('ignore')

# === CREDENTIAL SETUP WITH VALIDATION ===
def setup_credentials():
    """Setup and validate IBM Watson credentials - simplified sequential steps"""
    import os
    from getpass import getpass
    
    # Step 1: Try os.getenv first (checks if already in environment)
    print("🔍 Step 1: Checking environment variables...")
    ibm_api_key = os.getenv('IBM_API_KEY')
    ibm_project_id = os.getenv('IBM_PROJECT_ID')
    watsonx_apikey = os.getenv('WATSONX_APIKEY')
    
    if ibm_api_key and ibm_project_id:
        print("✅ Found credentials in environment variables")
        if not watsonx_apikey:
            watsonx_apikey = ibm_api_key
    else:
        # Step 2: Try dotenv to load .env file
        print("🔍 Step 2: Trying to load .env file...")
        try:
            from dotenv import load_dotenv
            from pathlib import Path
            
            # Try home directory first, then find_dotenv
            env_path = Path.home() / '.env'
            if env_path.exists():
                load_dotenv(env_path)
                print(f"📄 Loaded .env from {env_path}")
            else:
                from dotenv import find_dotenv
                env_file = find_dotenv()
                if env_file:
                    load_dotenv(env_file)
                    print(f"📄 Loaded .env from {env_file}")
            
            # Check environment again after loading .env
            ibm_api_key = os.getenv('IBM_API_KEY')
            ibm_project_id = os.getenv('IBM_PROJECT_ID')
            watsonx_apikey = os.getenv('WATSONX_APIKEY')
            
            if ibm_api_key and ibm_project_id:
                print("✅ Found credentials from .env file")
                if not watsonx_apikey:
                    watsonx_apikey = ibm_api_key
            else:
                print("⚠️  .env file loaded but missing required credentials")
                
        except ImportError:
            print("⚠️  python-dotenv not installed")
        except Exception as e:
            print(f"⚠️  Error loading .env: {e}")
    
    # Step 3: Manual input if still missing
    if not ibm_api_key or not ibm_project_id:
        print("🔍 Step 3: Manual input required...")
        ibm_api_key = getpass('IBM API Key: ')
        ibm_project_id = getpass('IBM Project ID: ')
        watsonx_apikey = ibm_api_key
    
    # Step 4: Validate and set
    if not ibm_api_key or len(ibm_api_key) < 40:
        raise ValueError("IBM API Key appears invalid (too short)")
    if not ibm_project_id or len(ibm_project_id) < 30:
        raise ValueError("IBM Project ID appears invalid (too short)")
    
    # Set environment variables
    os.environ['IBM_API_KEY'] = ibm_api_key
    os.environ['IBM_PROJECT_ID'] = ibm_project_id
    os.environ['WATSONX_APIKEY'] = watsonx_apikey
    
    print("✅ Credentials configured successfully")
    return ibm_api_key, ibm_project_id, watsonx_apikey

# === TEST CREDENTIALS FUNCTION ===
def test_credentials():
    """Test if credentials work by making a simple API call"""
    try:
        from ibm_watsonx_ai.foundation_models import ModelInference
        from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
        
        # Try to create a simple model instance
        model = ModelInference(
            model_id='mistralai/mixtral-8x7b-instruct-v01',
            params={GenParams.MAX_NEW_TOKENS: 10},
            credentials={
                "url": "https://us-south.ml.cloud.ibm.com",
                "apikey": os.environ['IBM_API_KEY']
            },
            project_id=os.environ['IBM_PROJECT_ID']
        )
        
        # Try a simple generation
        response = model.generate_text("Hello")
        print("✅ Credentials test successful!")
        return True
        
    except Exception as e:
        print(f"❌ Credentials test failed: {e}")
        return False

# === IMPORTS ===
from ibm_watsonx_ai.foundation_models import ModelInference
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames as GenParams
from ibm_watsonx_ai.metanames import EmbedTextParamsMetaNames
from langchain_ibm import WatsonxLLM, WatsonxEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain.chains import RetrievalQA
import gradio as gr

# === CORE FUNCTIONS WITH ERROR HANDLING ===

def get_llm():
    """Initialize the LLM with error handling"""
    try:
        model_id = 'mistralai/mixtral-8x7b-instruct-v01'
        parameters = {
            GenParams.MAX_NEW_TOKENS: 256,
            GenParams.TEMPERATURE: 0.5,
        }
        
        watsonx_llm = WatsonxLLM(
            model_id=model_id,
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=parameters,
        )
        return watsonx_llm
    except Exception as e:
        print(f"❌ Error initializing LLM: {e}")
        raise

def watsonx_embedding():
    """Initialize embeddings with error handling"""
    try:
        embed_params = {
            EmbedTextParamsMetaNames.TRUNCATE_INPUT_TOKENS: 3,
            EmbedTextParamsMetaNames.RETURN_OPTIONS: {"input_text": True},
        }
        
        watsonx_embedding = WatsonxEmbeddings(
            model_id="ibm/slate-125m-english-rtrvr",
            url="https://us-south.ml.cloud.ibm.com",
            project_id=os.environ['IBM_PROJECT_ID'],
            apikey=os.environ['IBM_API_KEY'],
            params=embed_params,
        )
        return watsonx_embedding
    except Exception as e:
        print(f"❌ Error initializing embeddings: {e}")
        raise

def get_file_icon(file_extension):
    """Get appropriate emoji icon for file type"""
    icons = {
        '.pdf': '📄',
        '.doc': '📝',
        '.docx': '📝',
        '.txt': '📃'
    }
    return icons.get(file_extension.lower(), '📎')

def document_loader_universal(file):
    """
    Universal document loader supporting PDF, DOC, DOCX, and TXT files
    Enhanced with better error handling and DOC support
    """
    import os
    from pathlib import Path
    
    try:
        # Validate file input
        if not file or not hasattr(file, 'name'):
            raise ValueError("Invalid file provided")
        
        file_path = file.name
        if not os.path.exists(file_path):
            raise ValueError(f"File does not exist: {file_path}")
        
        # Get file extension and icon
        file_extension = Path(file_path).suffix.lower()
        file_icon = get_file_icon(file_extension)
        
        print(f"{file_icon} Loading {file_extension.upper()} file: {file_path}")
        
        # Route to appropriate loader based on file type
        if file_extension == '.pdf':
            loader = PyPDFLoader(file_path)
            
        elif file_extension in ['.doc', '.docx']:
            # Enhanced DOC/DOCX support
            try:
                loader = Docx2txtLoader(file_path)
            except Exception as docx_error:
                # Fallback for DOC files or corrupted DOCX
                print(f"⚠️  Docx2txtLoader failed: {docx_error}")
                try:
                    # Alternative approach using python-docx directly
                    from docx import Document
                    from langchain.schema import Document as LangChainDocument
                    
                    if file_extension == '.docx':
                        doc = Document(file_path)
                        full_text = []
                        for paragraph in doc.paragraphs:
                            full_text.append(paragraph.text)
                        
                        content = '\n'.join(full_text)
                        if not content.strip():
                            raise ValueError("DOCX file appears to be empty")
                        
                        # Create document manually
                        loaded_document = [LangChainDocument(
                            page_content=content,
                            metadata={'source': file_path, 'file_type': file_extension}
                        )]
                        
                        print(f"✅ Successfully loaded DOCX using fallback method")
                        return self._finalize_document_loading(loaded_document, file_path, file_extension)
                    else:
                        raise ValueError(".DOC files require conversion to .DOCX format")
                        
                except ImportError:
                    raise ValueError("python-docx package required for DOCX files. Install with: pip install python-docx")
                except Exception as fallback_error:
                    raise ValueError(f"Failed to load {file_extension} file: {fallback_error}")
            
        elif file_extension == '.txt':
            # Enhanced text file support with encoding detection
            encodings_to_try = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
            
            loader = None
            successful_encoding = None
            
            for encoding in encodings_to_try:
                try:
                    loader = TextLoader(file_path, encoding=encoding)
                    # Test if it can actually load
                    test_load = loader.load()
                    successful_encoding = encoding
                    print(f"✅ Successfully detected text encoding: {encoding}")
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
                except Exception:
                    continue
            
            if loader is None:
                raise ValueError(f"Could not determine text file encoding. Tried: {encodings_to_try}")
                
        else:
            supported_types = ['.pdf', '.doc', '.docx', '.txt']
            raise ValueError(f"Unsupported file type: {file_extension}. Supported types: {supported_types}")
        
        # Load the document
        loaded_document = loader.load()
        
        return self._finalize_document_loading(loaded_document, file_path, file_extension)
        
    except Exception as e:
        error_msg = f"Error loading document '{os.path.basename(file_path) if 'file_path' in locals() else 'unknown'}': {str(e)}"
        print(f"❌ {error_msg}")
        raise ValueError(error_msg)

def _finalize_document_loading(loaded_document, file_path, file_extension):
    """Helper function to finalize document loading with validation and metadata"""
    import os
    
    # Validate content was loaded
    if not loaded_document:
        raise ValueError(f"No content found in file: {os.path.basename(file_path)}")
    
    # Check if documents have content
    total_chars = sum(len(doc.page_content) for doc in loaded_document)
    if total_chars == 0:
        raise ValueError(f"File appears to be empty: {os.path.basename(file_path)}")
    
    file_icon = get_file_icon(file_extension)
    print(f"✅ Successfully loaded {len(loaded_document)} document section(s) with {total_chars:,} characters")
    
    # Enhanced metadata
    for i, doc in enumerate(loaded_document):
        doc.metadata.update({
            'file_type': file_extension,
            'source_file': os.path.basename(file_path),
            'file_size_bytes': os.path.getsize(file_path),
            'total_sections': len(loaded_document),
            'section_index': i,
            'character_count': len(doc.page_content),
            'file_icon': file_icon
        })
    
    return loaded_document

def text_splitter(data):
    """Split text into chunks with improved configuration"""
    try:
        if not data:
            raise ValueError("No document data provided")
        
        # Improved text splitter configuration
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,  # Increased overlap for better continuity
            length_function=len,
            separators=["\n\n", "\n", " ", ""]  # Better separation hierarchy
        )
        chunks = text_splitter.split_documents(data)
        
        if not chunks:
            raise ValueError("No chunks created from document")
        
        print(f"✅ Split document into {len(chunks)} chunks")
        return chunks
    except Exception as e:
        print(f"❌ Error splitting text: {e}")
        raise

def vector_database(chunks):
    """Create vector database with improved configuration"""
    try:
        if not chunks:
            raise ValueError("No chunks provided for vector database")
        
        embedding_model = watsonx_embedding()
        
        # Configure Chroma with improved settings
        vectordb = Chroma.from_documents(
            documents=chunks, 
            embedding=embedding_model,
            collection_metadata={"hnsw:space": "cosine"}  # Better similarity metric
        )
        print("✅ Created vector database")
        return vectordb
    except Exception as e:
        print(f"❌ Error creating vector database: {e}")
        raise

def create_retriever(file):
    """Create retriever from file with improved configuration"""
    try:
        splits = document_loader_universal(file)
        chunks = text_splitter(splits)
        vectordb = vector_database(chunks)
        
        # Configure retriever with dynamic k based on chunk count
        max_chunks = len(chunks)
        k = min(4, max_chunks)  # Use minimum of 4 or available chunks
        
        retriever = vectordb.as_retriever(
            search_type="similarity",
            search_kwargs={"k": k}  # Dynamic k value
        )
        print(f"✅ Created retriever with k={k} (max available: {max_chunks})")
        return retriever, splits[0].metadata if splits else {}
    except Exception as e:
        print(f"❌ Error creating retriever: {e}")
        raise

def retriever_qa(file, query):
    """Main QA function with comprehensive error handling and file type display"""
    try:
        # Input validation
        if not file:
            return "Please upload a document file first."
        
        if not query or not query.strip():
            return "Please enter a question."
        
        print(f"🔍 Processing query: {query}")
        print(f"📎 File: {file.name if hasattr(file, 'name') else file}")
        
        # Initialize components
        llm = get_llm()
        retriever_obj, file_metadata = create_retriever(file)
        
        # Create QA chain
        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever_obj,
            return_source_documents=False,
            verbose=False
        )
        
        # Use invoke instead of deprecated __call__ method
        try:
            response = qa.invoke({"query": query})
            result = response.get('result', response) if isinstance(response, dict) else response
        except AttributeError:
            # Fallback for older LangChain versions
            response = qa({"query": query})
            result = response['result']
        
        # Add file info to response
        file_icon = file_metadata.get('file_icon', '📎')
        source_file = file_metadata.get('source_file', 'document')
        file_type = file_metadata.get('file_type', '').upper()
        
        response_with_source = f"{result}\n\n---\n💡 *Source: {file_icon} {source_file} ({file_type})*"
        
        print("✅ Generated response successfully")
        return response_with_source
        
    except ValueError as ve:
        error_msg = f"Input Error: {str(ve)}"
        print(f"❌ {error_msg}")
        return error_msg
        
    except Exception as e:
        error_msg = f"Processing Error: {str(e)}"
        print(f"❌ {error_msg}")
        return f"Sorry, I encountered an error while processing your request.\n\nError details: {str(e)}\n\nPlease check your file and try again."

# === ENHANCED GRADIO INTERFACE ===
def create_gradio_interface():
    """Create an enhanced Gradio interface with multi-format support"""
    
    def get_file_info(file):
        """Extract and display file information"""
        if not file:
            return "No file uploaded"
        
        try:
            import os
            from pathlib import Path
            
            file_path = file.name if hasattr(file, 'name') else str(file)
            file_name = os.path.basename(file_path)
            file_extension = Path(file_path).suffix.lower()
            file_size = os.path.getsize(file_path)
            file_icon = get_file_icon(file_extension)
            
            # Format file size
            if file_size < 1024:
                size_str = f"{file_size} bytes"
            elif file_size < 1024*1024:
                size_str = f"{file_size/1024:.1f} KB"
            else:
                size_str = f"{file_size/(1024*1024):.1f} MB"
            
            return f"{file_icon} **{file_name}** ({file_extension.upper()}, {size_str})"
            
        except Exception as e:
            return f"Error reading file info: {str(e)}"
    
    def process_query(file, query, history, file_info):
        """Process query and maintain chat history"""
        if not file:
            response = "⚠️ Please upload a document file first."
            file_info_text = "No file uploaded"
        elif not query.strip():
            response = "⚠️ Please enter a question."
            file_info_text = get_file_info(file)
        else:
            response = retriever_qa(file, query)
            file_info_text = get_file_info(file)
        
        # Add to history
        history.append([query, response])
        return history, "", file_info_text
    
    def update_file_info(file):
        """Update file info display when file is uploaded"""
        return get_file_info(file)
    
    def clear_chat():
        """Clear chat history"""
        return [], "No file uploaded"
    
    # Create the interface
    with gr.Blocks(title="RAG Document Q&A System", theme=gr.themes.Soft()) as app:
        gr.Markdown("# 🤖 RAG Document Q&A System")
        gr.Markdown("Upload your document (PDF, DOC, DOCX, or TXT) and ask questions about its content!")
        
        with gr.Row():
            with gr.Column(scale=1):
                file_upload = gr.File(
                    label="📎 Upload Document", 
                    file_count="single", 
                    file_types=['.pdf', '.doc', '.docx', '.txt'],
                    type="filepath",
                    height=100
                )
                
                file_info_display = gr.Markdown(
                    value="No file uploaded",
                    label="File Information"
                )
                
                with gr.Row():
                    clear_btn = gr.Button("🗑️ Clear Chat", variant="secondary", scale=1)
                
                gr.Markdown("""
                ### 📋 Supported Formats:
                - **📄 PDF** - Portable Document Format
                - **📝 DOC/DOCX** - Microsoft Word documents  
                - **📃 TXT** - Plain text files
                
                ### 💡 Tips:
                - Ask specific questions about the content
                - Try questions like "What is the main topic?" or "Summarize the key points"
                - For resumes: "What skills are mentioned?" or "What is the work experience?"
                """)
                
            with gr.Column(scale=2):
                chatbot = gr.Chatbot(
                    label="💬 Chat History",
                    height=500,
                    show_label=True,
                    avatar_images=("👤", "🤖")
                )
                
                with gr.Row():
                    query_input = gr.Textbox(
                        label="Your Question",
                        placeholder="Ask a question about the uploaded document...",
                        lines=2,
                        scale=4
                    )
                    submit_btn = gr.Button("📤 Send", variant="primary", scale=1, size="lg")
        
        # Event handlers
        file_upload.upload(
            update_file_info,
            inputs=[file_upload],
            outputs=[file_info_display]
        )
        
        submit_btn.click(
            process_query,
            inputs=[file_upload, query_input, chatbot, file_info_display],
            outputs=[chatbot, query_input, file_info_display]
        )
        
        query_input.submit(
            process_query,
            inputs=[file_upload, query_input, chatbot, file_info_display],
            outputs=[chatbot, query_input, file_info_display]
        )
        
        clear_btn.click(
            clear_chat,
            outputs=[chatbot, file_info_display]
        )
    
    return app

# === MAIN SETUP ===
def main():
    """Main setup function with enhanced error handling"""
    print("🚀 Setting up Multi-Format RAG Application...")
    
    try:
        # Setup credentials
        setup_credentials()
        
        # Test credentials
        if not test_credentials():
            print("❌ Credential test failed. Please check your IBM Watson credentials.")
            return None
        
        # Create Gradio interface
        app = create_gradio_interface()
        
        print("✅ Multi-Format RAG Application ready!")
        print("📋 Supported formats: PDF, DOC, DOCX, TXT")
        return app
        
    except Exception as e:
        print(f"❌ Failed to initialize application: {e}")
        return None

# === LAUNCH ===
if __name__ == "__main__":
    # Setup and launch
    app = main()
    if app:
        print("🌐 Launching application...")
        app.launch(
            server_name="0.0.0.0", 
            server_port=7860, 
            share=False, 
            debug=False,
            show_error=True
        )
    else:
        print("❌ Failed to initialize application")

# === REQUIREMENTS ===
"""
Required packages (install with pip):

pip install python-docx          # For DOC/DOCX support
pip install python-dotenv        # For environment variables
pip install langchain-community  # For document loaders
pip install langchain-ibm        # For Watson integration
pip install chromadb            # For vector database
pip install gradio              # For web interface
pip install ibm-watsonx-ai      # For Watson AI
"""

