from docx import Document
from docx.shared import Inches
import google.generativeai as genai

# Read document with images
doc = Document('input_with_images.docx')

# Extract text (images are ignored)
text_content = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():  # Only non-empty paragraphs
        text_content.append(paragraph.text)

# Send text to Gemini for rewriting
model = genai.GenerativeModel('gemini-pro')
rewritten_text = model.generate_content(f"Rewrite: {' '.join(text_content)}").text

# Create new document
new_doc = Document()
new_doc.add_paragraph(rewritten_text)

# You'd need to manually re-add images if needed
# new_doc.add_picture('image.jpg', width=Inches(4))

new_doc.save('rewritten.docx')

