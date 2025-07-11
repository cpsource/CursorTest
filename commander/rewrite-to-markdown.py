#!/usr/bin/env python3
"""
DOCX to Markdown Converter
Converts .docx files to markdown format, preserving structure and handling images.
"""

import os
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

class DocxToMarkdownConverter:
    def __init__(self, docx_path, extract_images=True):
        """Initialize converter with .docx file path"""
        self.docx_path = Path(docx_path)
        self.extract_images = extract_images
        self.output_dir = self.docx_path.parent
        self.images_dir = self.output_dir / f"{self.docx_path.stem}_images"
        self.markdown_file = self.output_dir / f"{self.docx_path.stem}.md"
        
    def extract_images_from_docx(self):
        """Extract images from .docx file"""
        if not self.extract_images:
            return {}
        
        self.images_dir.mkdir(exist_ok=True)
        image_mapping = {}
        
        try:
            with zipfile.ZipFile(self.docx_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/')]
                
                for img_file in image_files:
                    img_data = docx_zip.read(img_file)
                    img_name = os.path.basename(img_file)
                    
                    # Save image
                    output_path = self.images_dir / img_name
                    with open(output_path, 'wb') as f:
                        f.write(img_data)
                    
                    # Map internal reference to filename
                    image_mapping[img_file] = img_name
                    
        except Exception as e:
            print(f"Warning: Could not extract images: {e}")
            
        return image_mapping
    
    def get_paragraph_style(self, paragraph):
        """Determine markdown formatting for paragraph"""
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or 'title' in style_name:
            return '# '
        elif 'heading 2' in style_name:
            return '## '
        elif 'heading 3' in style_name:
            return '### '
        elif 'heading 4' in style_name:
            return '#### '
        elif 'heading 5' in style_name:
            return '##### '
        elif 'heading 6' in style_name:
            return '###### '
        elif 'quote' in style_name or 'blockquote' in style_name:
            return '> '
        else:
            return ''
    
    def format_run_text(self, run):
        """Apply markdown formatting to run text"""
        text = run.text
        
        if not text.strip():
            return text
            
        # Apply formatting
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        
        if run.underline:
            text = f"<u>{text}</u>"
            
        return text
    
    def process_paragraph(self, paragraph, image_mapping):
        """Convert paragraph to markdown"""
        # Check if paragraph contains images
        has_image = any(run.element.xpath('.//a:blip') for run in paragraph.runs)
        
        if has_image and self.extract_images:
            # Add image placeholder
            image_text = f"\n![Image](./{self.images_dir.name}/image_placeholder.png)\n"
            return image_text
        
        # Get paragraph style
        style_prefix = self.get_paragraph_style(paragraph)
        
        # Process runs (formatted text segments)
        formatted_text = ""
        for run in paragraph.runs:
            formatted_text += self.format_run_text(run)
        
        # Handle empty paragraphs
        if not formatted_text.strip():
            return "\n"
        
        # Apply paragraph-level formatting
        result = style_prefix + formatted_text
        
        # Handle alignment (for special cases)
        if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            result = f"<div align='center'>{formatted_text}</div>"
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            result = f"<div align='right'>{formatted_text}</div>"
        
        return result + "\n"
    
    def process_table(self, table):
        """Convert table to markdown"""
        markdown_table = []
        
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                # Get cell text, handling multiple paragraphs
                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                # Clean up cell text for markdown
                cell_text = cell_text.replace('\n', ' ').replace('|', '\\|')
                row_cells.append(cell_text)
            
            # Add table row
            markdown_table.append("| " + " | ".join(row_cells) + " |")
            
            # Add header separator after first row
            if row_idx == 0:
                separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                markdown_table.append(separator)
        
        return "\n".join(markdown_table) + "\n\n"
    
    def convert_to_markdown(self):
        """Main conversion method"""
        try:
            # Extract images first
            image_mapping = self.extract_images_from_docx()
            
            # Load document
            doc = Document(self.docx_path)
            markdown_content = []
            
            # Process document elements
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    for paragraph in doc.paragraphs:
                        if paragraph.element == element:
                            md_text = self.process_paragraph(paragraph, image_mapping)
                            markdown_content.append(md_text)
                            break
                            
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    for table in doc.tables:
                        if table.element == element:
                            md_table = self.process_table(table)
                            markdown_content.append(md_table)
                            break
            
            # Write markdown file
            final_content = "".join(markdown_content)
            
            # Clean up extra newlines
            final_content = re.sub(r'\n{3,}', '\n\n', final_content)
            
            with open(self.markdown_file, 'w', encoding='utf-8') as f:
                f.write(final_content)
            
            return {
                'markdown_file': self.markdown_file,
                'images_extracted': len(image_mapping),
                'images_directory': self.images_dir if image_mapping else None
            }
            
        except Exception as e:
            raise Exception(f"Conversion failed: {e}")

def main():
    """Example usage"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docx_to_markdown.py <docx_file> [--no-images]")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    extract_images = '--no-images' not in sys.argv
    
    try:
        converter = DocxToMarkdownConverter(docx_file, extract_images)
        result = converter.convert_to_markdown()
        
        print(f"✅ Conversion complete!")
        print(f"📄 Markdown file: {result['markdown_file']}")
        if result['images_extracted'] > 0:
            print(f"🖼️  Images extracted: {result['images_extracted']}")
            print(f"📁 Images directory: {result['images_directory']}")
        
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()
