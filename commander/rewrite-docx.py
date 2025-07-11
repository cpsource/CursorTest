from docx import Document
import google.generativeai as genai

# Configure Gemini
genai.configure(api_key="your-api-key")
model = genai.GenerativeModel('gemini-pro')

# Read the .docx file
doc = Document('input.docx')
original_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

# Send to Gemini for rewriting
prompt = f"Rewrite this text to be more professional: {original_text}"
response = model.generate_content(prompt)
rewritten_text = response.text

# Create new document with rewritten content
new_doc = Document()
new_doc.add_paragraph(rewritten_text)
new_doc.save('rewritten.docx')

