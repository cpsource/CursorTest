#!/usr/bin/env python3
"""
DOCX Image Extractor Tool
Extracts all embedded images from a .docx file and saves them to a directory.
"""

import os
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import shutil

class DocxImageExtractor:
    def __init__(self, docx_path):
        """Initialize with path to .docx file"""
        self.docx_path = Path(docx_path)
        self.output_dir = self.docx_path.parent / f"{self.docx_path.stem}_images"
        
    def extract_images(self):
        """Extract all images from the .docx file"""
        if not self.docx_path.exists():
            raise FileNotFoundError(f"File not found: {self.docx_path}")
        
        # Create output directory
        self.output_dir.mkdir(exist_ok=True)
        
        # .docx files are actually zip files
        with zipfile.ZipFile(self.docx_path, 'r') as docx_zip:
            # Get list of image files in the media folder
            image_files = [f for f in docx_zip.namelist() 
                          if f.startswith('word/media/')]
            
            if not image_files:
                print("No images found in the document.")
                return []
            
            extracted_images = []
            
            for img_file in image_files:
                # Extract the image
                img_data = docx_zip.read(img_file)
                
                # Get just the filename (remove path)
                img_name = os.path.basename(img_file)
                
                # Save to output directory
                output_path = self.output_dir / img_name
                with open(output_path, 'wb') as f:
                    f.write(img_data)
                
                extracted_images.append(output_path)
                print(f"Extracted: {img_name}")
            
            return extracted_images
    
    def get_image_info(self):
        """Get information about images and their positions in the document"""
        doc = Document(self.docx_path)
        image_info = []
        
        # Check paragraphs for images
        for i, paragraph in enumerate(doc.paragraphs):
            # Look for drawing elements (images) in the paragraph
            for run in paragraph.runs:
                if run.element.xpath('.//a:blip'):
                    # Found an image in this paragraph
                    image_info.append({
                        'paragraph_index': i,
                        'paragraph_text': paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text,
                        'has_image': True
                    })
                    break
        
        # Check tables for images
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.element.xpath('.//a:blip'):
                                image_info.append({
                                    'location': f'Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}',
                                    'paragraph_text': paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text,
                                    'has_image': True
                                })
        
        return image_info
    
    def extract_with_context(self):
        """Extract images and return them with context about where they appear"""
        extracted_images = self.extract_images()
        image_info = self.get_image_info()
        
        return {
            'extracted_files': extracted_images,
            'image_locations': image_info,
            'output_directory': self.output_dir
        }

def main():
    """Example usage"""
    import sys
    
    if len(sys.argv) != 2:
        print("Usage: python docx_image_extractor.py <path_to_docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    try:
        extractor = DocxImageExtractor(docx_file)
        result = extractor.extract_with_context()
        
        print(f"\n--- Extraction Summary ---")
        print(f"Output directory: {result['output_directory']}")
        print(f"Total images extracted: {len(result['extracted_files'])}")
        
        if result['image_locations']:
            print("\n--- Image Locations ---")
            for location in result['image_locations']:
                if 'paragraph_index' in location:
                    print(f"Paragraph {location['paragraph_index']}: {location['paragraph_text']}")
                else:
                    print(f"{location['location']}: {location['paragraph_text']}")
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()

